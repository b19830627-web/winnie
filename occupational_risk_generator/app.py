import io
import json
from datetime import datetime
from pathlib import Path

import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


APP_TITLE = "特約職護職業安全風險評估產生器"
RULES_PATH = Path(__file__).with_name("risk_rules.json")
DEFAULT_ACCESS_PASSWORD = "666"


FIELD_LABELS = {
    "department": "部門",
    "job_title": "職務名稱",
    "work_content": "主要作業內容",
    "shift": "班別",
    "health_risks": "特殊健康風險",
    "health_management_year": "健康管理年度",
    "health_level": "健康管理分級",
    "health_case_id": "健康編號",
    "health_basis": "健康評估依據",
    "senior_assessment": "中高齡評估",
    "work_ability_level": "中高齡工作適能等級",
    "medical_follow_up": "評估結果",
    "special_hazard_year": "特別危害作業年度",
    "work_fitness_options": "工作適性建議",
    "work_fitness_detail": "工作適性詳述",
    "notes": "補充說明",
}


WORK_ABILITY_MEANINGS = {
    "弱": "不能勝任工作要求",
    "普通": "工作適能有待提高",
    "良": "能勝任所從事的工作",
    "優": "能很好地勝任所從事的工作",
}

MEDICAL_FOLLOW_UP_OPTIONS = [
    "正常或不需要治療，可結案。",
    "已診斷治療中，可結案。",
    "已複診檢查及追蹤，可結案。",
    "需進一步檢查及持續追蹤(項目)",
    "自主管理，並於年度健檢追蹤。",
    "需立即就醫治療",
]

WORK_FITNESS_OPTIONS = [
    "適任原單位工作",
    "進行工作調整",
    "工作禁忌",
    "工作限制",
    "不適任",
]

HEALTH_LEVEL_OPTIONS = ["", "第一級", "第二級", "第三級", "第四級"]

SPECIAL_HAZARD_FIELDS = {
    "noise_health_level": "噪音",
    "dust_health_level": "粉塵",
    "radiation_health_level": "游離輻射",
}

HEALTH_STATUS_KEYWORDS = ["不影響工作", "已不影響", "無功能受限", "已恢復", "已復原"]


def get_access_password() -> str:
    try:
        return st.secrets.get("APP_PASSWORD", DEFAULT_ACCESS_PASSWORD)
    except Exception:
        return DEFAULT_ACCESS_PASSWORD


def require_password() -> bool:
    if st.session_state.get("authenticated"):
        with st.sidebar:
            st.success("已登入")
            if st.button("登出"):
                st.session_state["authenticated"] = False
                st.rerun()
        return True

    st.title(APP_TITLE)
    st.caption("請先輸入使用密碼。")
    password = st.text_input("密碼", type="password")
    if st.button("登入", type="primary"):
        if password == get_access_password():
            st.session_state["authenticated"] = True
            st.rerun()
        else:
            st.error("密碼錯誤，請重新輸入。")
    return False


WORKPLACE_KEYWORDS = [
    "作業",
    "工作",
    "現場",
    "環境",
    "人因",
    "姿勢",
    "動線",
    "設備",
    "輔具",
    "防護具",
    "通風",
    "休息",
    "工時",
    "排班",
    "輪班",
    "疲勞",
    "搬運",
    "重量",
    "頻率",
    "暴露",
    "高溫",
    "噪音",
    "粉塵",
    "化學",
    "交通",
    "改善",
    "通報",
]

MEDICAL_FOCUS_KEYWORDS = [
    "就醫",
    "回診",
    "服藥",
    "用藥",
    "醫療",
    "治療",
    "醫師面談",
    "檢查結果",
    "健檢",
    "健康管理",
    "健康追蹤",
    "血壓",
    "血糖",
    "血脂",
]

HEALTH_LABEL_KEYWORDS = ["高血壓", "血糖", "血脂", "肝功能", "貧血", "中高齡", "慢性病", "孕產婦", "健康管理"]


def load_rules() -> dict:
    with RULES_PATH.open("r", encoding="utf-8") as file:
        return json.load(file)


def normalize_text(value) -> str:
    if isinstance(value, list):
        return "、".join(str(item).strip() for item in value if str(item).strip())
    return str(value or "").strip()


def field_contains(form_data: dict, fields: list[str], keywords: list[str]) -> bool:
    for field in fields:
        value = normalize_text(form_data.get(field))
        if any(keyword in value for keyword in keywords):
            return True
    return False


def dedupe(items: list[str]) -> list[str]:
    seen = set()
    result = []
    for item in items:
        clean = item.strip()
        if clean and clean not in seen:
            seen.add(clean)
            result.append(clean)
    return result


def get_matched_rules(form_data: dict, rules_data: dict) -> list[dict]:
    matched = []
    for rule in rules_data["rules"]:
        match = rule.get("match", {})
        if field_contains(form_data, match.get("fields", []), match.get("keywords", [])):
            matched.append(rule)
    return matched


def get_context_notes(form_data: dict, rules_data: dict) -> list[str]:
    notes = []
    note_map = rules_data.get("industry_department_notes", {})
    for key in [form_data.get("department")]:
        clean_key = normalize_text(key)
        if clean_key in note_map:
            notes.extend(note_map[clean_key])
    return dedupe(notes)


def build_intro(form_data: dict) -> str:
    parts = []
    for key in ["department", "job_title"]:
        value = normalize_text(form_data.get(key))
        if value:
            parts.append(f"{FIELD_LABELS[key]}：{value}")
    return "；".join(parts)


def build_section(title: str, items: list[str]) -> str:
    if not items:
        return f"{title}\n- 尚無可產生之具體內容。"
    lines = [title]
    lines.extend(f"- {item}" for item in items)
    return "\n".join(lines)


def build_numbered_lines(items: list[str]) -> list[str]:
    return [f"{index}. {item}" for index, item in enumerate(items, start=1)]


def split_custom_health_risks(custom_risks: str) -> list[str]:
    return [
        item.strip()
        for item in custom_risks.replace("，", ",").replace("、", ",").split(",")
        if item.strip()
    ]


def merge_health_risks(selected_risks: list[str], custom_risks: str) -> list[str]:
    custom_items = [
        item for item in split_custom_health_risks(custom_risks)
        if not any(keyword in item for keyword in HEALTH_STATUS_KEYWORDS)
    ]
    merged = dedupe(selected_risks + custom_items)
    if "無" in merged and len(merged) > 1:
        merged = [risk for risk in merged if risk != "無"]
    return merged or ["無"]


def get_custom_health_risk_guidance(custom_risks: list[str], work_content: str, shift: str) -> dict[str, list[str]]:
    event_risks = [
        item for item in custom_risks
        if not any(keyword in item for keyword in HEALTH_STATUS_KEYWORDS)
    ]
    if not event_risks:
        return {"risk": [], "education": [], "improvements": [], "follow_up": []}

    risk_text = "、".join(event_risks)
    work_text = normalize_text(work_content) or "現行作業"
    shift_text = normalize_text(shift)
    management_items = [
        f"針對其他健康風險註記（{risk_text}），建議檢視{work_text}之工作負荷、作業頻率、休息安排及高風險作業分派",
        f"若個案近期有{risk_text}相關病史或功能受限情形，建議確認疼痛、不適或功能受限程度及作業安全適配性",
    ]
    if shift_text and shift_text not in ["日班", "未填"]:
        management_items.append(f"針對{shift_text}作業，建議同步檢視工時安排、連續工作時間及疲勞累積情形")

    return {
        "risk": [
            f"其他健康風險註記為{risk_text}，需結合實際作業姿勢、體力負荷、移動需求、暴露因子及班別安排，評估是否增加職場傷病、跌倒、疲勞或作業安全事件風險",
        ],
        "education": [
            "已指導個案於作業時留意疼痛、發燒、頭暈、疲倦、步態不穩、傷口或感染相關警訊，若工作中出現明顯不適，應暫停高風險作業並通報主管或職護",
            "提醒個案依作業內容採取適當休息、伸展、補水及防護具使用，避免因健康狀況未穩定而增加職業傷害風險",
        ],
        "improvements": management_items,
        "follow_up": [
            f"建議於下次臨場服務追蹤其他健康風險（{risk_text}）對{work_text}之作業負荷、工作適性及現場改善措施執行情形",
        ],
    }


def get_special_hazard_guidance(form_data: dict) -> dict[str, list[str]]:
    guidance = {"management": [], "environment": [], "education": [], "follow_up": []}
    noise_level = normalize_text(form_data.get("noise_health_level"))
    dust_level = normalize_text(form_data.get("dust_health_level"))
    radiation_level = normalize_text(form_data.get("radiation_health_level"))

    if noise_level:
        guidance["management"].append(
            f"噪音作業管理：目前噪音作業健康管理為{noise_level}，建議檢視噪音暴露時間、作業輪替、聽力防護具配戴及保存管理情形。"
        )
        guidance["environment"].append(
            "噪音作業環境建議：確認噪音源隔離、設備維護、警示標示及耳塞或耳罩配置是否符合現場作業需求。"
        )
        guidance["education"].append(
            "噪音作業教育指導：提醒正確配戴耳塞或耳罩，避免自行縮短配戴時間，並留意耳鳴、聽力下降或溝通困難等警訊。"
        )
        guidance["follow_up"].append(
            "後續追蹤噪音暴露控制、聽力防護具使用情形及噪音作業健康管理分級變化。"
        )

    if dust_level:
        guidance["management"].append(
            f"粉塵作業管理：目前粉塵作業健康管理為{dust_level}，建議檢視粉塵暴露時間、作業方式、局部排氣及呼吸防護具使用管理。"
        )
        guidance["environment"].append(
            "粉塵作業環境建議：確認濕式作業、局部排氣、清掃方式、粉塵逸散控制及呼吸防護具配置是否足夠。"
        )
        guidance["education"].append(
            "粉塵作業教育指導：提醒正確配戴合適呼吸防護具，避免乾掃或造成二次揚塵，並留意咳嗽、胸悶、呼吸道刺激等症狀。"
        )
        guidance["follow_up"].append(
            "後續追蹤粉塵暴露控制、呼吸防護具使用情形及粉塵作業健康管理分級變化。"
        )

    if radiation_level:
        guidance["management"].append(
            f"游離輻射作業管理：目前游離輻射作業健康管理為{radiation_level}，建議檢視作業許可、暴露時間、劑量佩章配戴、屏蔽措施及作業紀錄管理。"
        )
        guidance["environment"].append(
            "游離輻射作業環境建議：確認管制區標示、屏蔽設備、警示裝置、劑量監測與人員進出管理是否落實。"
        )
        guidance["education"].append(
            "游離輻射作業教育指導：提醒遵守時間、距離、屏蔽原則，正確配戴劑量佩章，並依規定完成作業紀錄與異常通報。"
        )
        guidance["follow_up"].append(
            "後續追蹤游離輻射暴露紀錄、劑量佩章配戴情形及游離輻射作業健康管理分級變化。"
        )

    return guidance


def has_special_hazard_level(form_data: dict) -> bool:
    return any(normalize_text(form_data.get(field)) for field in SPECIAL_HAZARD_FIELDS)


def join_as_record_sentence(items: list[str]) -> str:
    cleaned = []
    for item in items:
        clean = item.strip().rstrip("。；;")
        if clean:
            cleaned.append(clean)
    if not cleaned:
        return ""
    return "；".join(cleaned) + "。"


def strip_record_punctuation(text: str) -> str:
    return text.strip().rstrip("。；;")


def normalize_year(year: str) -> str:
    clean_year = normalize_text(year)
    return clean_year.rstrip("年度年") if clean_year else ""


def with_year_prefix(text: str, year: str) -> str:
    clean_year = normalize_year(year)
    if clean_year:
        return f"{clean_year}年度{text}"
    return text


def build_data_basis(health_level: str, form_data: dict) -> str:
    health_year = normalize_year(form_data.get("health_management_year"))
    special_hazard_year = normalize_year(form_data.get("special_hazard_year")) or health_year

    if health_level and health_level != "未分級":
        base_level_text = health_level if health_level.startswith("健康管理") else f"健康管理{health_level}"
        level_text = with_year_prefix(base_level_text, health_year)
    else:
        level_text = "未提供健康管理分級"

    special_hazard_parts = []
    for field, label in SPECIAL_HAZARD_FIELDS.items():
        level = normalize_text(form_data.get(field))
        if level:
            special_hazard_parts.append(with_year_prefix(f"{label}作業健康管理{level}", special_hazard_year))

    if special_hazard_parts:
        return f"{level_text}；特別危害作業健康管理：{'、'.join(special_hazard_parts)}。"
    return f"{level_text}。"


def build_risk_focus(health_risks: str, work_risks: str, form_data: dict) -> str:
    health_categories = [
        item.strip()
        for item in health_risks.replace("，", "、").split("、")
        if item.strip() and item.strip() not in ["無", "無特殊健康風險註記"]
    ]
    hazard_categories = [
        f"{label}作業暴露"
        for field, label in SPECIAL_HAZARD_FIELDS.items()
        if normalize_text(form_data.get(field))
    ]
    main_categories = dedupe(health_categories + hazard_categories)

    parts = []
    if main_categories:
        parts.append(f"主要風險包括{'、'.join(main_categories)}。")

    exposure_text = normalize_text(work_risks)
    if exposure_text and exposure_text != "未明確辨識":
        source_text = f"{exposure_text} {normalize_text(form_data.get('work_content'))}"
        outcome_map = [
            (["久站", "久坐", "搬運", "重複", "人因", "姿勢", "手部", "電腦"], "肌肉骨骼傷害"),
            (["噪音"], "聽力損失"),
            (["粉塵", "化學品", "呼吸"], "呼吸系統不適"),
            (["輪班", "夜間", "長工時", "疲勞", "高溫", "熱"], "心血管與疲勞負荷"),
            (["高溫", "熱"], "熱危害"),
            (["游離輻射", "輻射"], "游離輻射暴露危害"),
            (["外勤", "交通", "駕駛", "重機具"], "作業事故"),
            (["監督", "工程進度", "品質管理", "安全管理", "人員協調"], "工作壓力、疲勞及判斷失誤"),
        ]
        outcomes = dedupe([
            outcome
            for keywords, outcome in outcome_map
            if any(keyword in source_text for keyword in keywords)
        ])
        if outcomes:
            parts.append(f"持續暴露於{exposure_text}，可能增加{'、'.join(outcomes)}風險。")

    work_ability_level = normalize_text(form_data.get("work_ability_level"))
    if work_ability_level:
        ability_summaries = {
            "弱": "目前尚無法勝任現行工作要求，應進一步評估工作調整需求",
            "普通": "目前工作適能仍有待提高，應持續評估作業負荷與適配情形",
            "良": "目前具備勝任現行工作的能力",
            "優": "目前具備良好勝任現行工作的能力",
        }
        parts.append(
            f"中高齡工作適能指數為「{work_ability_level}」，"
            f"{ability_summaries.get(work_ability_level, WORK_ABILITY_MEANINGS.get(work_ability_level, ''))}。"
        )

    senior_assessment = normalize_text(form_data.get("senior_assessment"))
    if senior_assessment:
        parts.append(f"{strip_record_punctuation(senior_assessment)}。")

    return "".join(parts) or "目前未辨識出明確風險，建議補充作業內容、班別或暴露因子，以利完整評估。"


def keep_workplace_item(item: str) -> bool:
    return any(keyword in item for keyword in WORKPLACE_KEYWORDS) and not any(keyword in item for keyword in MEDICAL_FOCUS_KEYWORDS)


def normalize_role_wording(item: str) -> str:
    replacements = {
        "建議雇主與職護共同評估是否需調整工作負荷、班別、暴露時間或高風險作業安排": "建議檢視工作負荷、班別、暴露時間及高風險作業安排是否需調整",
        "建議主管與職護共同評估工作負荷、休息安排與必要之工作調整": "建議檢視工作負荷、休息安排及必要之工作調整",
        "建議雇主依醫師與職護建議": "建議依評估結果",
        "建議主管與職護共同": "建議",
        "主管與職護共同": "",
        "建議主管依": "建議依",
        "建議主管": "建議",
        "建議雇主": "建議",
    }
    normalized = item
    for old, new in replacements.items():
        normalized = normalized.replace(old, new)
    return normalized


def extract_recent_health_context(custom_risks: list[str]) -> str:
    event_risks = [
        item for item in custom_risks
        if not any(keyword in item for keyword in HEALTH_STATUS_KEYWORDS)
    ]
    if not event_risks:
        return ""
    risk_text = "、".join(event_risks)
    if "因" in risk_text:
        before, after = risk_text.split("因", 1)
        date_text = before.strip()
        reason_text = after.strip()
        if date_text and reason_text:
            return f"個案近期（{date_text}）因{reason_text}"
    return f"個案近期有{risk_text}相關健康風險"


def build_workplace_management(items: list[str], work_risks: str, shift: str, custom_risks=None) -> str:
    custom_risks = custom_risks or []
    has_heat_risk = any(
        keyword in work_risks
        for keyword in ["高溫作業", "高氣溫作業", "熱暴露", "熱環境", "熱危害"]
    )
    grouped_items = []
    if has_heat_risk:
        grouped_items.append(
            "熱危害與負荷管理：建立熱危害通報與緊急處置流程；依現場熱暴露程度與高氣溫作業特性，適時檢視並調整作業分派、輪替制度及休息頻率。"
        )
    if custom_risks:
        health_context = extract_recent_health_context(custom_risks)
        has_resolved_status = any(
            any(keyword in item for keyword in HEALTH_STATUS_KEYWORDS) for item in custom_risks
        )
        if health_context and has_resolved_status:
            grouped_items.append(
                f"傷病適配管理：{health_context}，經評估已不影響工作，現階段無須針對此傷病調整工作內容。"
            )
        elif health_context:
            grouped_items.append(
                f"傷病適配管理：{health_context}，建議確認目前症狀、功能受限程度及作業安全適配性，必要時調整工作內容。"
            )
    if grouped_items:
        for item in items:
            clean_item = normalize_role_wording(strip_record_punctuation(item))
            if clean_item.startswith(("噪音作業管理", "粉塵作業管理", "游離輻射作業管理", "工作負荷與壓力管理")):
                grouped_items.append(f"{clean_item}。")
        return "\n".join(dedupe(grouped_items))

    cleaned_items = [
        normalize_role_wording(strip_record_punctuation(item)) for item in items if keep_workplace_item(item)
    ]
    if work_risks and work_risks != "未明確辨識":
        cleaned_items.append(f"建議依{work_risks}檢視作業分派、休息頻率、輪替制度與工作負荷")
    if shift and shift not in ["未填", "日班"]:
        cleaned_items.append(f"針對{shift}人員建立疲勞風險辨識、交接班確認與異常狀況通報機制")
    return join_as_record_sentence(dedupe(cleaned_items))


def build_workplace_education(items: list[str], work_risks: str) -> str:
    cleaned_items = [strip_record_punctuation(item) for item in items if keep_workplace_item(item)]
    if work_risks and work_risks != "未明確辨識":
        cleaned_items.append(f"教育個案辨識{work_risks}相關作業危害，工作中若出現明顯不適、注意力下降或動作控制受影響，應暫停高風險作業並通報主管或職護")
    cleaned_items.append("指導依實際作業姿勢執行伸展、補水、休息與防護具正確使用，以降低職業傷害風險")
    return join_as_record_sentence(dedupe(cleaned_items))


def build_workplace_follow_up(follow_up_items: list[str], work_risks: str, work_content: str, shift: str) -> str:
    cleaned_items = []
    for item in follow_up_items:
        clean = strip_record_punctuation(item)
        if not clean:
            continue
        if keep_workplace_item(clean):
            cleaned_items.append(clean)

    if work_risks and work_risks != "未明確辨識":
        cleaned_items.append(f"建議於下次臨場服務追蹤{work_risks}相關作業安排、環境改善與人因工程調整成效")
    if work_content and work_content != "未填":
        cleaned_items.append(f"持續確認{work_content}之實際作業負荷、動線配置及休息安排是否符合現場需求")
    if shift and shift not in ["未填", "日班"]:
        cleaned_items.append(f"追蹤{shift}之工時安排、疲勞累積與交接班制度對作業安全之影響")

    return join_as_record_sentence(dedupe(cleaned_items))


def generate_recommendation(form_data: dict, rules_data: dict) -> tuple[str, list[dict]]:
    matched_rules = get_matched_rules(form_data, rules_data)
    context_notes = get_context_notes(form_data, rules_data)
    has_exposure_data = bool(
        normalize_text(form_data.get("work_content"))
        or normalize_text(form_data.get("shift"))
        or normalize_text(form_data.get("health_level"))
        or has_special_hazard_level(form_data)
    )
    has_only_none_health = form_data.get("health_risks") == ["無"]

    if not matched_rules and not context_notes and (not has_exposure_data or has_only_none_health):
        message = rules_data["insufficient_data_message"]
        intro = build_intro(form_data)
        text = "\n\n".join(
            part
            for part in [
                APP_TITLE,
                intro,
                build_section("一、職業安全衛生風險評估", [message]),
                build_section("二、衛教建議", [message]),
                build_section("三、具體改善措施", [message]),
                build_section("四、後續追蹤建議", [message]),
            ]
            if part
        )
        return text, matched_rules

    risk_items = []
    education_items = []
    improvement_items = []
    follow_up_items = []

    risk_items.extend(context_notes)
    for rule in matched_rules:
        risk_items.extend(rule.get("risk", []))
        education_items.extend(rule.get("education", []))
        improvement_items.extend(rule.get("improvements", []))
        follow_up_items.extend(rule.get("follow_up", []))

    custom_guidance = get_custom_health_risk_guidance(
        form_data.get("custom_health_risks", []),
        form_data.get("work_content", ""),
        form_data.get("shift", ""),
    )
    special_hazard_guidance = get_special_hazard_guidance(form_data)
    risk_items.extend(custom_guidance["risk"])
    education_items.extend(custom_guidance["education"])
    improvement_items.extend(custom_guidance["improvements"])
    follow_up_items.extend(custom_guidance["follow_up"])
    education_items.extend(special_hazard_guidance["education"])
    improvement_items.extend(special_hazard_guidance["management"])
    improvement_items.extend(special_hazard_guidance["environment"])
    follow_up_items.extend(special_hazard_guidance["follow_up"])

    if not risk_items:
        risk_items.append(rules_data["insufficient_data_message"])
    if not education_items:
        education_items.append("建議補充主要作業內容與暴露因子後，再提供更精準之衛教內容。")
    if not improvement_items:
        improvement_items.append("建議補充現場作業流程、設備配置與暴露情形後，再提出具體改善措施。")

    follow_up_items.extend(rules_data.get("default_follow_up", []))

    notes = normalize_text(form_data.get("notes"))
    if notes:
        risk_items.append(f"補充說明顯示：{notes}。建議職護於臨場服務時進一步確認其與作業暴露及健康風險之關聯。")

    intro = build_intro(form_data)
    work_summary = []
    for key in ["work_content", "shift", "health_risks", "health_level"]:
        value = normalize_text(form_data.get(key))
        if value:
            work_summary.append(f"{FIELD_LABELS[key]}：{value}")

    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
    document_parts = [
        APP_TITLE,
        f"產出時間：{generated_at}",
        intro,
        "；".join(work_summary),
        build_section("一、職業安全衛生風險評估", dedupe(risk_items)),
        build_section("二、衛教建議", dedupe(education_items)),
        build_section("三、具體改善措施", dedupe(improvement_items)),
        build_section("四、後續追蹤建議", dedupe(follow_up_items)),
    ]
    return "\n\n".join(part for part in document_parts if part), matched_rules


def generate_appendix_record(form_data: dict, rules_data: dict) -> tuple[str, list[dict]]:
    matched_rules = get_matched_rules(form_data, rules_data)
    context_notes = get_context_notes(form_data, rules_data)
    custom_health_risks = form_data.get("custom_health_risks", [])

    if not matched_rules and not context_notes and not custom_health_risks and not has_special_hazard_level(form_data):
        message = rules_data["insufficient_data_message"]
        return "\n\n".join(
            part
            for part in [
                "1.健康編號：未填",
                f"(1)基本資料與作業特性\n1-1 工作型態：{message}\n1-2 作業特性：{message}",
                f"(2)健康評估與風險分析\n2-1 數據依據：{message}\n2-2 風險重點：{message}",
                f"(3)改善及建議採行措施\n3-1 管理建議：{message}\n3-2 環境建議：{message}\n3-3 教育指導：{message}",
                f"(4)適性評估與後續追蹤\n4-1 健康指導結果：{message}\n4-2 工作適性評估：{message}\n4-3 後續追蹤：{message}",
            ]
            if part
        ), matched_rules

    risk_items = []
    education_items = []
    improvement_items = []
    follow_up_items = []
    risk_items.extend(context_notes)
    for rule in matched_rules:
        risk_items.extend(rule.get("risk", []))
        education_items.extend(rule.get("education", []))
        improvement_items.extend(rule.get("improvements", []))
        follow_up_items.extend(rule.get("follow_up", []))

    custom_guidance = get_custom_health_risk_guidance(
        form_data.get("custom_health_risks", []),
        form_data.get("work_content", ""),
        form_data.get("shift", ""),
    )
    special_hazard_guidance = get_special_hazard_guidance(form_data)
    risk_items.extend(custom_guidance["risk"])
    education_items.extend(custom_guidance["education"])
    improvement_items.extend(custom_guidance["improvements"])
    follow_up_items.extend(custom_guidance["follow_up"])
    education_items.extend(special_hazard_guidance["education"])
    improvement_items.extend(special_hazard_guidance["management"])
    improvement_items.extend(special_hazard_guidance["environment"])
    follow_up_items.extend(special_hazard_guidance["follow_up"])

    risk_items = dedupe(risk_items)
    education_items = dedupe(education_items)
    improvement_items = dedupe(improvement_items)
    follow_up_items = dedupe(follow_up_items)

    case_id = normalize_text(form_data.get("health_case_id")) or "未填"
    department = normalize_text(form_data.get("department")) or "未填"
    job_title = normalize_text(form_data.get("job_title")) or "未填"
    shift = normalize_text(form_data.get("shift")) or "未填"
    work_content = normalize_text(form_data.get("work_content")) or "未填"
    health_risks = normalize_text(form_data.get("health_risks")) or "無特殊健康風險註記"
    health_level = normalize_text(form_data.get("health_level")) or "未分級"
    senior_assessment = normalize_text(form_data.get("senior_assessment"))
    work_ability_level = normalize_text(form_data.get("work_ability_level"))
    medical_follow_up = normalize_text(form_data.get("medical_follow_up"))
    work_fitness_options = normalize_text(form_data.get("work_fitness_options"))
    work_fitness_detail = normalize_text(form_data.get("work_fitness_detail"))
    notes = normalize_text(form_data.get("notes"))

    management_items = []
    environment_items = []
    for item in improvement_items:
        if item.startswith(("噪音作業管理", "粉塵作業管理", "游離輻射作業管理")):
            management_items.append(item)
        elif item.startswith(("噪音作業環境建議", "粉塵作業環境建議", "游離輻射作業環境建議")):
            environment_items.append(item)
        elif any(keyword in item for keyword in ["耳塞", "耳罩", "配戴方式", "保存狀況", "防護具尺寸"]):
            management_items.append(item)
        elif any(keyword in item for keyword in ["工作檯", "物品", "設備", "通風", "遮陽", "補水點", "防護具", "推車", "升降台", "動線", "座椅", "螢幕", "鍵盤", "滑鼠", "環境", "配置"]):
            environment_items.append(item)
        else:
            management_items.append(item)

    if not management_items:
        management_items.append("建議將個案納入健康管理追蹤名冊，定期掌握健康檢查結果，並評估其與現行作業之適配情形。")
    if not environment_items:
        environment_items.append("建議實地檢視作業空間、設備配置、動線與人因工程條件，必要時調整以降低作業負荷。")
    if not education_items:
        education_items.append("建議由職護提供個別健康指導，說明作業相關危害、症狀警訊與自我保護方式。")

    if notes:
        clean_notes = notes.rstrip("。；;")
        risk_items.append(f"補充說明：{clean_notes}。建議於臨場服務時進一步確認其與作業負荷及健康風險之關聯。")

    matched_labels = [rule["label"] for rule in matched_rules]
    exposure_labels = [
        label for label in matched_labels if not any(keyword in label for keyword in HEALTH_LABEL_KEYWORDS)
    ]
    all_risks = "、".join(matched_labels) or "未明確辨識"
    work_risks = "、".join(exposure_labels) or (work_content if work_content != "未填" else "未明確辨識")
    data_basis = build_data_basis(health_level, form_data)
    risk_focus = build_risk_focus(health_risks, work_risks, form_data)
    management_text = build_workplace_management(
        management_items,
        work_risks,
        shift,
        form_data.get("custom_health_risks", []),
    )
    environment_text = join_as_record_sentence([normalize_role_wording(item) for item in environment_items])
    education_text = build_workplace_education(education_items, work_risks)
    follow_up_text = build_workplace_follow_up(follow_up_items, work_risks, work_content, shift)
    fit_result = normalize_text(form_data.get("fit_result"))
    if not fit_result:
        selected_fitness = work_fitness_options or f"適任原單位工作"
        fit_result = f"現階段評估為{selected_fitness}。"
        if work_fitness_detail:
            fit_result = f"{fit_result}{strip_record_punctuation(work_fitness_detail)}。"
        elif "適任原單位工作" in selected_fitness:
            fit_result = f"{fit_result}可勝任{department}{job_title}及原作業內容。"
        else:
            fit_result = f"{fit_result}建議依職護、醫師與事業單位評估結果，確認必要之工作調整、限制或後續追蹤事項。"
    if medical_follow_up:
        fit_result = f"{fit_result} 評估結果：{strip_record_punctuation(medical_follow_up)}。"
    health_guidance_result = f"已向個案說明{all_risks}對健康與作業安全之影響，並給予與作業風險相關之健康指導及改善建議。"
    management_line = f"3-1 管理建議：\n{management_text}" if "\n" in management_text else f"3-1 管理建議：{management_text}"
    case_heading = f"{case_id}（{department}）" if department != "未填" else case_id

    parts = [
        f"1.健康編號：{case_heading}",
        "(1)基本資料與作業特性",
        f"1-1 工作型態：工作班別為{shift}。",
        f"1-2 作業特性：擔任{job_title}，主要作業內容為{strip_record_punctuation(work_content)}。",
        "",
        "(2)健康評估與風險分析",
        f"2-1 數據依據：{data_basis}",
        f"2-2 風險重點：{risk_focus}",
        "",
        "(3)改善及建議採行措施",
        management_line,
        f"3-2 環境建議： {environment_text}",
        f"3-3 教育指導： {education_text}",
        "",
        "(4)適性評估與後續追蹤",
        f"4-1 健康指導結果：{health_guidance_result}",
        f"4-2 工作適性評估：{fit_result}",
        f"4-3 後續追蹤：{follow_up_text}",
    ]

    return "\n".join(parts), matched_rules


def create_word_file(text: str, form_data: dict) -> bytes:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Microsoft JhengHei"
    styles["Normal"].font.size = Pt(11)

    lines = [line.rstrip() for line in text.splitlines()]
    title_text = next((line for line in lines if line.strip()), APP_TITLE)

    title = document.add_heading(title_text, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    skip_first_title = True
    for line in lines:
        clean_line = line.strip()
        if not clean_line:
            continue
        if skip_first_title and clean_line == title_text:
            skip_first_title = False
            continue

        if clean_line.startswith(("一、", "二、", "三、", "四、", "(1)", "(2)", "(3)", "(4)")):
            document.add_heading(clean_line, level=2)
        elif clean_line.startswith(("3-1", "3-2", "3-3")):
            document.add_heading(clean_line, level=3)
        elif clean_line.startswith("- "):
            document.add_paragraph(clean_line.removeprefix("- ").strip(), style="List Bullet")
        elif len(clean_line) > 3 and clean_line[0].isdigit() and clean_line[1:3] == ". ":
            document.add_paragraph(clean_line[3:].strip(), style="List Number")
        else:
            document.add_paragraph(clean_line)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def copy_button(text: str) -> None:
    escaped_text = json.dumps(text, ensure_ascii=False)
    components.html(
        f"""
        <button id="copyBtn" style="
            background:#2563eb;color:white;border:0;border-radius:6px;
            padding:0.55rem 0.9rem;font-size:0.95rem;cursor:pointer;">
            複製完整文字
        </button>
        <span id="copyStatus" style="margin-left:0.75rem;color:#166534;font-size:0.9rem;"></span>
        <script>
        const btn = document.getElementById("copyBtn");
        const status = document.getElementById("copyStatus");
        btn.onclick = async () => {{
            try {{
                await navigator.clipboard.writeText({escaped_text});
                status.textContent = "已複製";
            }} catch (error) {{
                status.textContent = "複製失敗，請手動選取文字";
            }}
        }};
        </script>
        """,
        height=48,
    )


def main() -> None:
    st.set_page_config(page_title=APP_TITLE, page_icon="🩺", layout="wide")
    if not require_password():
        return

    rules_data = load_rules()

    st.title(APP_TITLE)
    st.caption("依據部門、職務、作業內容、班別與健康風險，自動產生可用於臨場服務紀錄、年度報告或健康風險評估表之文字建議。")

    with st.form("risk_form"):
        output_format = st.radio(
            "輸出格式",
            ["四大段建議", "附表八紀錄式"],
            horizontal=True,
            help="附表八紀錄式會產生基本資料、健康評估與風險分析、管理／環境／教育改善措施、適工評估與後續追蹤。",
        )
        col1, col2 = st.columns(2)
        with col1:
            department = st.text_input("部門", placeholder="例如：行政部、倉庫、製造部、品保、工程部、房務、櫃檯、廚房、照服員")
            job_title = st.text_input("職務名稱")
            health_risks = st.multiselect(
                "特殊健康風險",
                ["高血壓", "血糖異常", "血脂異常", "肝功能異常", "貧血", "肌肉骨骼不適", "疲勞高負荷", "中高齡", "孕產婦", "慢性病", "無"],
                default=["無"],
            )
            custom_health_risks = st.text_input(
                "其他健康風險",
                placeholder="可自行輸入，例如：住院、手術後、睡眠障礙；多個項目可用頓號或逗號分隔",
            )
        with col2:
            shift = st.selectbox("班別", ["", "日班", "二班制", "三班制", "輪班", "夜間工作"])
            health_management_year = st.text_input(
                "健康管理年度",
                placeholder="例如：114、115",
                help="填入民國年度數字，輸出會顯示為「115年度健康管理第二級」。",
            )
            health_level = st.selectbox(
                "健康管理分級",
                HEALTH_LEVEL_OPTIONS,
                help="可依公司健檢管理或職護評估結果選擇健康管理分級。",
            )
            work_content = st.text_area(
                "主要作業內容",
                placeholder="例如：久站、久坐、搬運、重複性手部操作、電腦作業、外勤移動、輪班、夜間工作、高溫作業、噪音、粉塵、化學品接觸",
                height=110,
            )
            notes = st.text_area("補充說明", height=80)

        with st.expander("附表八紀錄欄位（選填）"):
            col3, col4 = st.columns(2)
            with col3:
                health_case_id = st.text_input("健康編號", placeholder="例如：115-04-15-01")
                health_basis = st.text_area("健康評估依據", placeholder="例如：113 年度健康檢查第三級管理、血壓第二級管理。", height=80)
                st.markdown("特別危害作業健康管理")
                special_hazard_year = st.text_input(
                    "特別危害作業年度",
                    placeholder="例如：114、115；未填則沿用健康管理年度",
                )
                noise_health_level = st.selectbox("噪音", HEALTH_LEVEL_OPTIONS)
                dust_health_level = st.selectbox("粉塵", HEALTH_LEVEL_OPTIONS)
                radiation_health_level = st.selectbox("游離輻射", HEALTH_LEVEL_OPTIONS)
                work_ability_level = st.selectbox(
                    "中高齡工作適能等級",
                    ["", "弱", "普通", "良", "優"],
                    help="弱：不能勝任工作要求；普通：工作適能有待提高；良：能勝任所從事的工作；優：能很好地勝任所從事的工作。",
                )
                senior_assessment = st.text_area("中高齡評估補充", placeholder="例如：視力、聽力、認知、肌力皆正常。", height=70)
            with col4:
                medical_follow_up = st.multiselect("評估結果", MEDICAL_FOLLOW_UP_OPTIONS)
                work_fitness_options = st.multiselect("工作適性建議", WORK_FITNESS_OPTIONS, default=["適任原單位工作"])
                work_fitness_detail = st.text_area("工作適性詳述", placeholder="例如：可適任原職務工作；若症狀惡化，建議再行評估作業適配情形。", height=70)
                fit_result = st.text_area("工作適性評估自訂文字", placeholder="若填寫此欄，會優先使用此文字。", height=70)

        submitted = st.form_submit_button("產生建議", type="primary")

    if submitted:
        custom_health_risk_items = split_custom_health_risks(custom_health_risks)
        health_risks = merge_health_risks(health_risks, custom_health_risks)

        form_data = {
            "company": "",
            "industry": "",
            "department": department,
            "job_title": job_title,
            "work_content": work_content,
            "shift": shift,
            "health_risks": health_risks,
            "custom_health_risks": custom_health_risk_items,
            "health_management_year": health_management_year,
            "health_level": health_level,
            "health_case_id": health_case_id,
            "health_basis": health_basis,
            "special_hazard_year": special_hazard_year,
            "noise_health_level": noise_health_level,
            "dust_health_level": dust_health_level,
            "radiation_health_level": radiation_health_level,
            "senior_assessment": senior_assessment,
            "work_ability_level": work_ability_level,
            "medical_follow_up": medical_follow_up,
            "work_fitness_options": work_fitness_options,
            "work_fitness_detail": work_fitness_detail,
            "fit_result": fit_result,
            "notes": notes,
        }
        if output_format == "附表八紀錄式":
            result_text, matched_rules = generate_appendix_record(form_data, rules_data)
        else:
            result_text, matched_rules = generate_recommendation(form_data, rules_data)
        st.session_state["result_text"] = result_text
        st.session_state["form_data"] = form_data
        st.session_state["matched_rules"] = [rule["label"] for rule in matched_rules]
        st.session_state["output_format"] = output_format

    if "result_text" in st.session_state:
        st.subheader("產生結果")
        if st.session_state.get("matched_rules"):
            st.info("已套用規則：" + "、".join(st.session_state["matched_rules"]))
        else:
            st.warning(rules_data["insufficient_data_message"])

        st.text_area("完整文字", st.session_state["result_text"], height=520)
        button_col1, button_col2 = st.columns([1, 2])
        with button_col1:
            copy_button(st.session_state["result_text"])
        with button_col2:
            docx_bytes = create_word_file(st.session_state["result_text"], st.session_state["form_data"])
            safe_department = normalize_text(st.session_state["form_data"].get("department")) or "職業安全風險評估"
            suffix = "附表八紀錄" if st.session_state.get("output_format") == "附表八紀錄式" else "職業安全風險評估建議"
            st.download_button(
                "下載 Word 檔",
                data=docx_bytes,
                file_name=f"{safe_department}_{suffix}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    with st.expander("規則庫維護說明"):
        st.write(
            "本程式會讀取同資料夾的 `risk_rules.json`。日後可新增 `rules` 內的規則，"
            "設定 `keywords`、`risk`、`education`、`improvements` 與 `follow_up` 後即可擴充產出內容。"
        )


if __name__ == "__main__":
    main()
