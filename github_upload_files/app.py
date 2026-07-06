import io
import json
from datetime import datetime
from pathlib import Path

import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


APP_TITLE = "特約職護職業安全風險評估產生器"
RULES_PATH = Path(__file__).with_name("risk_rules.json")


FIELD_LABELS = {
    "company": "公司名稱",
    "industry": "產業別",
    "department": "部門",
    "job_title": "職務名稱",
    "work_content": "主要作業內容",
    "shift": "班別",
    "health_risks": "特殊健康風險",
    "health_level": "健康管理分級",
    "health_case_id": "健康編號",
    "health_basis": "健康評估依據",
    "senior_assessment": "中高齡評估",
    "notes": "補充說明",
}


def load_rules() -> dict:
    with RULES_PATH.open("r", encoding="utf-8") as file:
        return json.load(file)


def normalize_text(value) -> str:
    if isinstance(value, list):
        return "、".join(str(item).strip() for item in value if str(item).strip())
    return str(value or "").strip()


def field_contains(form_data: dict, fields: list[str], keywords: list[str]) -> bool:
    for field in fields:
        value = normalize_text(form_data.get(field))
        if any(keyword in value for keyword in keywords):
            return True
    return False


def dedupe(items: list[str]) -> list[str]:
    seen = set()
    result = []
    for item in items:
        clean = item.strip()
        if clean and clean not in seen:
            seen.add(clean)
            result.append(clean)
    return result


def get_matched_rules(form_data: dict, rules_data: dict) -> list[dict]:
    matched = []
    for rule in rules_data["rules"]:
        match = rule.get("match", {})
        if field_contains(form_data, match.get("fields", []), match.get("keywords", [])):
            matched.append(rule)
    return matched


def get_context_notes(form_data: dict, rules_data: dict) -> list[str]:
    notes = []
    note_map = rules_data.get("industry_department_notes", {})
    for key in [form_data.get("industry"), form_data.get("department")]:
        clean_key = normalize_text(key)
        if clean_key in note_map:
            notes.extend(note_map[clean_key])
    return dedupe(notes)


def build_intro(form_data: dict) -> str:
    parts = []
    for key in ["company", "industry", "department", "job_title"]:
        value = normalize_text(form_data.get(key))
        if value:
            parts.append(f"{FIELD_LABELS[key]}：{value}")
    return "；".join(parts)


def build_section(title: str, items: list[str]) -> str:
    if not items:
        return f"{title}\n- 尚無可產生之具體內容。"
    lines = [title]
    lines.extend(f"- {item}" for item in items)
    return "\n".join(lines)


def build_numbered_lines(items: list[str]) -> list[str]:
    return [f"{index}. {item}" for index, item in enumerate(items, start=1)]


def generate_recommendation(form_data: dict, rules_data: dict) -> tuple[str, list[dict]]:
    matched_rules = get_matched_rules(form_data, rules_data)
    context_notes = get_context_notes(form_data, rules_data)
    has_exposure_data = bool(
        normalize_text(form_data.get("work_content"))
        or normalize_text(form_data.get("shift"))
        or normalize_text(form_data.get("health_level"))
    )
    has_only_none_health = form_data.get("health_risks") == ["無"]

    if not matched_rules and not context_notes and (not has_exposure_data or has_only_none_health):
        message = rules_data["insufficient_data_message"]
        intro = build_intro(form_data)
        text = "\n\n".join(
            part
            for part in [
                APP_TITLE,
                intro,
                build_section("一、職業安全衛生風險評估", [message]),
                build_section("二、衛教建議", [message]),
                build_section("三、具體改善措施", [message]),
                build_section("四、後續追蹤建議", [message]),
            ]
            if part
        )
        return text, matched_rules

    risk_items = []
    education_items = []
    improvement_items = []
    follow_up_items = []

    risk_items.extend(context_notes)
    for rule in matched_rules:
        risk_items.extend(rule.get("risk", []))
        education_items.extend(rule.get("education", []))
        improvement_items.extend(rule.get("improvements", []))
        follow_up_items.extend(rule.get("follow_up", []))

    if not risk_items:
        risk_items.append(rules_data["insufficient_data_message"])
    if not education_items:
        education_items.append("建議補充主要作業內容與暴露因子後，再提供更精準之衛教內容。")
    if not improvement_items:
        improvement_items.append("建議補充現場作業流程、設備配置與暴露情形後，再提出具體改善措施。")

    follow_up_items.extend(rules_data.get("default_follow_up", []))

    notes = normalize_text(form_data.get("notes"))
    if notes:
        risk_items.append(f"補充說明顯示：{notes}。建議職護於臨場服務時進一步確認其與作業暴露及健康風險之關聯。")

    intro = build_intro(form_data)
    work_summary = []
    for key in ["work_content", "shift", "health_risks", "health_level"]:
        value = normalize_text(form_data.get(key))
        if value:
            work_summary.append(f"{FIELD_LABELS[key]}：{value}")

    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
    document_parts = [
        APP_TITLE,
        f"產出時間：{generated_at}",
        intro,
        "；".join(work_summary),
        build_section("一、職業安全衛生風險評估", dedupe(risk_items)),
        build_section("二、衛教建議", dedupe(education_items)),
        build_section("三、具體改善措施", dedupe(improvement_items)),
        build_section("四、後續追蹤建議", dedupe(follow_up_items)),
    ]
    return "\n\n".join(part for part in document_parts if part), matched_rules


def generate_appendix_record(form_data: dict, rules_data: dict) -> tuple[str, list[dict]]:
    matched_rules = get_matched_rules(form_data, rules_data)
    context_notes = get_context_notes(form_data, rules_data)

    if not matched_rules and not context_notes:
        message = rules_data["insufficient_data_message"]
        intro = build_intro(form_data)
        return "\n\n".join(
            part
            for part in [
                "附表八健康風險評估與適工追蹤紀錄",
                intro,
                f"(1)基本資料與作業特性\n{message}",
                f"(2)健康評估與風險分析\n{message}",
                f"(3)改善建議與採行措施\n{message}",
                f"(4)適工評估與後續追蹤\n{message}",
            ]
            if part
        ), matched_rules

    risk_items = []
    education_items = []
    improvement_items = []
    follow_up_items = []
    risk_items.extend(context_notes)
    for rule in matched_rules:
        risk_items.extend(rule.get("risk", []))
        education_items.extend(rule.get("education", []))
        improvement_items.extend(rule.get("improvements", []))
        follow_up_items.extend(rule.get("follow_up", []))

    risk_items = dedupe(risk_items)
    education_items = dedupe(education_items)
    improvement_items = dedupe(improvement_items)
    follow_up_items = dedupe(follow_up_items + rules_data.get("default_follow_up", []))

    case_id = normalize_text(form_data.get("health_case_id")) or "未填"
    company = normalize_text(form_data.get("company")) or "未填"
    department = normalize_text(form_data.get("department")) or "未填"
    job_title = normalize_text(form_data.get("job_title")) or "未填"
    industry = normalize_text(form_data.get("industry")) or "未填"
    shift = normalize_text(form_data.get("shift")) or "未填"
    work_content = normalize_text(form_data.get("work_content")) or "未填"
    health_risks = normalize_text(form_data.get("health_risks")) or "無特殊健康風險註記"
    health_level = normalize_text(form_data.get("health_level")) or "未分級"
    health_basis = normalize_text(form_data.get("health_basis"))
    senior_assessment = normalize_text(form_data.get("senior_assessment"))
    notes = normalize_text(form_data.get("notes"))

    basis_text = health_basis or f"依健康檢查結果及職護評估，目前健康管理分級為{health_level}，特殊健康風險註記為：{health_risks}。"
    senior_text = senior_assessment or "如屬中高齡勞工，建議依實際工作適能、視力、聽力、肌力、認知與慢性病控制情形補充評估。"

    management_items = []
    environment_items = []
    for item in improvement_items:
        if any(keyword in item for keyword in ["工作檯", "物品", "設備", "通風", "遮陽", "補水點", "防護具", "推車", "升降台", "動線", "座椅", "螢幕", "鍵盤", "滑鼠", "環境", "配置"]):
            environment_items.append(item)
        else:
            management_items.append(item)

    if not management_items:
        management_items.append("建議將個案納入健康管理追蹤名冊，定期掌握健康檢查結果，並評估其與現行作業之適配情形。")
    if not environment_items:
        environment_items.append("建議實地檢視作業空間、設備配置、動線與人因工程條件，必要時調整以降低作業負荷。")
    if not education_items:
        education_items.append("建議由職護提供個別健康指導，說明作業相關危害、症狀警訊與自我保護方式。")

    if notes:
        risk_items.append(f"補充說明：{notes}。建議於臨場服務時進一步確認其與作業負荷及健康風險之關聯。")

    work_risks = "、".join(rule["label"] for rule in matched_rules) or "未明確辨識"
    fit_result = normalize_text(form_data.get("fit_result")) or f"經本次評估，其健康狀況原則上可配合目前{shift}之{department}{job_title}作業；惟仍應依健康管理分級、症狀變化及現場暴露情形持續追蹤，必要時再行評估工作調整需求。"

    parts = [
        "附表八健康風險評估與適工追蹤紀錄",
        f"健康編號：{case_id}",
        "(1)基本資料與作業特性",
        f"公司名稱：{company}",
        f"產業別：{industry}",
        f"部門／職務：{department}／{job_title}",
        f"工作型態：{shift}。",
        f"作業特性評估：主要作業內容包含{work_content}。依輸入資料辨識之潛在風險包含：{work_risks}。",
        "",
        "(2)健康評估與風險分析",
        f"健康評估依據：{basis_text}",
        f"中高齡評估：{senior_text}",
        "風險分析：",
        *[f"- {item}" for item in risk_items],
        f"綜合評估：個案健康風險應與其作業型態、班別、暴露因子及健康管理分級一併判斷，並作為後續適工評估與健康管理追蹤之依據。",
        "",
        "(3)改善建議與採行措施",
        "針對上述風險，分別從管理、環境、教育三大面向提出具體對策：",
        "3-1 管理建議",
        *build_numbered_lines(management_items),
        "3-2 環境建議",
        *build_numbered_lines(environment_items),
        "3-3 教育指導",
        *build_numbered_lines(education_items),
        "",
        "(4)適工評估與後續追蹤",
        f"(1)適工性評估結果：{fit_result}",
        "(2)後續建議：",
        *[f"- {item}" for item in follow_up_items],
    ]

    return "\n".join(parts), matched_rules


def create_word_file(text: str, form_data: dict) -> bytes:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Microsoft JhengHei"
    styles["Normal"].font.size = Pt(11)

    lines = [line.rstrip() for line in text.splitlines()]
    title_text = next((line for line in lines if line.strip()), APP_TITLE)

    title = document.add_heading(title_text, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    skip_first_title = True
    for line in lines:
        clean_line = line.strip()
        if not clean_line:
            continue
        if skip_first_title and clean_line == title_text:
            skip_first_title = False
            continue

        if clean_line.startswith(("一、", "二、", "三、", "四、", "(1)", "(2)", "(3)", "(4)")):
            document.add_heading(clean_line, level=2)
        elif clean_line.startswith(("3-1", "3-2", "3-3")):
            document.add_heading(clean_line, level=3)
        elif clean_line.startswith("- "):
            document.add_paragraph(clean_line.removeprefix("- ").strip(), style="List Bullet")
        elif len(clean_line) > 3 and clean_line[0].isdigit() and clean_line[1:3] == ". ":
            document.add_paragraph(clean_line[3:].strip(), style="List Number")
        else:
            document.add_paragraph(clean_line)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def copy_button(text: str) -> None:
    escaped_text = json.dumps(text, ensure_ascii=False)
    components.html(
        f"""
        <button id="copyBtn" style="
            background:#2563eb;color:white;border:0;border-radius:6px;
            padding:0.55rem 0.9rem;font-size:0.95rem;cursor:pointer;">
            複製完整文字
        </button>
        <span id="copyStatus" style="margin-left:0.75rem;color:#166534;font-size:0.9rem;"></span>
        <script>
        const btn = document.getElementById("copyBtn");
        const status = document.getElementById("copyStatus");
        btn.onclick = async () => {{
            try {{
                await navigator.clipboard.writeText({escaped_text});
                status.textContent = "已複製";
            }} catch (error) {{
                status.textContent = "複製失敗，請手動選取文字";
            }}
        }};
        </script>
        """,
        height=48,
    )


def main() -> None:
    st.set_page_config(page_title=APP_TITLE, page_icon="🩺", layout="wide")
    rules_data = load_rules()

    st.title(APP_TITLE)
    st.caption("依據公司、部門、作業內容、班別與健康風險，自動產生可用於臨場服務紀錄、年度報告或健康風險評估表之文字建議。")

    with st.form("risk_form"):
        output_format = st.radio(
            "輸出格式",
            ["四大段建議", "附表八紀錄式"],
            horizontal=True,
            help="附表八紀錄式會產生基本資料、健康評估與風險分析、管理／環境／教育改善措施、適工評估與後續追蹤。",
        )
        col1, col2 = st.columns(2)
        with col1:
            company = st.text_input("公司名稱")
            industry = st.selectbox(
                "產業別",
                ["", "營造業", "製造業", "飯店業", "日照中心", "辦公室作業", "其他"],
            )
            department = st.text_input("部門", placeholder="例如：行政部、倉庫、製造部、品保、工程部、房務、櫃檯、廚房、照服員")
            job_title = st.text_input("職務名稱")
        with col2:
            shift = st.selectbox("班別", ["", "日班", "二班制", "三班制", "輪班", "夜間工作"])
            health_risks = st.multiselect(
                "是否有特殊健康風險",
                ["高血壓", "血糖異常", "血脂異常", "肝功能異常", "貧血", "肌肉骨骼不適", "疲勞高負荷", "中高齡", "孕產婦", "慢性病", "無"],
                default=["無"],
            )
            health_level = st.selectbox(
                "健康管理分級",
                ["", "第一級", "第二級", "第三級", "第四級"],
                help="可依公司健檢管理或職護評估結果選擇健康管理分級。",
            )
            work_content = st.text_area(
                "主要作業內容",
                placeholder="例如：久站、久坐、搬運、重複性手部操作、電腦作業、外勤移動、輪班、夜間工作、高溫作業、噪音、粉塵、化學品接觸",
                height=110,
            )
            notes = st.text_area("補充說明", height=80)

        with st.expander("附表八紀錄欄位（選填）"):
            col3, col4 = st.columns(2)
            with col3:
                health_case_id = st.text_input("健康編號", placeholder="例如：115-04-15-01")
                health_basis = st.text_area("健康評估依據", placeholder="例如：113 年度健康檢查第三級管理、血壓第二級管理。", height=80)
            with col4:
                senior_assessment = st.text_area("中高齡評估", placeholder="例如：工作適能指數優；視力、聽力、認知、肌力皆正常。", height=80)
                fit_result = st.text_area("適工性評估結果", placeholder="未填時會由系統依輸入內容產生一般性適工文字。", height=80)

        submitted = st.form_submit_button("產生建議", type="primary")

    if submitted:
        if "無" in health_risks and len(health_risks) > 1:
            health_risks = [risk for risk in health_risks if risk != "無"]

        form_data = {
            "company": company,
            "industry": industry,
            "department": department,
            "job_title": job_title,
            "work_content": work_content,
            "shift": shift,
            "health_risks": health_risks,
            "health_level": health_level,
            "health_case_id": health_case_id,
            "health_basis": health_basis,
            "senior_assessment": senior_assessment,
            "fit_result": fit_result,
            "notes": notes,
        }
        if output_format == "附表八紀錄式":
            result_text, matched_rules = generate_appendix_record(form_data, rules_data)
        else:
            result_text, matched_rules = generate_recommendation(form_data, rules_data)
        st.session_state["result_text"] = result_text
        st.session_state["form_data"] = form_data
        st.session_state["matched_rules"] = [rule["label"] for rule in matched_rules]
        st.session_state["output_format"] = output_format

    if "result_text" in st.session_state:
        st.subheader("產生結果")
        if st.session_state.get("matched_rules"):
            st.info("已套用規則：" + "、".join(st.session_state["matched_rules"]))
        else:
            st.warning(rules_data["insufficient_data_message"])

        st.text_area("完整文字", st.session_state["result_text"], height=520)
        button_col1, button_col2 = st.columns([1, 2])
        with button_col1:
            copy_button(st.session_state["result_text"])
        with button_col2:
            docx_bytes = create_word_file(st.session_state["result_text"], st.session_state["form_data"])
            safe_company = normalize_text(st.session_state["form_data"].get("company")) or "職業安全風險評估"
            suffix = "附表八紀錄" if st.session_state.get("output_format") == "附表八紀錄式" else "職業安全風險評估建議"
            st.download_button(
                "下載 Word 檔",
                data=docx_bytes,
                file_name=f"{safe_company}_{suffix}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    with st.expander("規則庫維護說明"):
        st.write(
            "本程式會讀取同資料夾的 `risk_rules.json`。日後可新增 `rules` 內的規則，"
            "設定 `keywords`、`risk`、`education`、`improvements` 與 `follow_up` 後即可擴充產出內容。"
        )


if __name__ == "__main__":
    main()
